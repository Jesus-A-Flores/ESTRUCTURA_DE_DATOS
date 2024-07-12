# Archivo: excel_to_word.py
from docx import Document
from docx.shared import Inches
import pandas as pd
import openpyxl
import os
from extract_charts import extract_charts_from_excel  # Importar función de extracción de gráficos

# Función para extraer imágenes de una hoja de Excel
def extract_images_from_excel(sheet, output_dir='images'):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    images = []
    for image in sheet._images:
        image_path = os.path.join(output_dir, f"{image.anchor._from.row}_{image.anchor._from.col}.png")
        image.ref.save(image_path)
        images.append((image.anchor._from.row, image.anchor._from.col, image_path))
    return images

if __name__ == "__main__":
    # Ruta del archivo de Excel
    excel_file = 'C:\\Users\\LENOVO\\Desktop\\CODIGOS\\pruevas\\Para_la_Vida\\RESULTADO.xlsx'  # Reemplaza con la ruta correcta

    # Leer el contenido del archivo de Excel
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb.active

    # Extraer imágenes de la hoja de Excel
    images = extract_images_from_excel(sheet)

    # Extraer gráficos de la hoja de Excel utilizando la función del otro archivo
    chart_images = extract_charts_from_excel(excel_file)

    # Leer datos de Excel en un DataFrame de pandas
    df = pd.read_excel(excel_file)

    # Crear un nuevo documento de Word
    doc = Document()

    # Agregar el contenido del DataFrame al documento de Word
    for i, row in df.iterrows():
        # Combinar valores de fila en una sola cadena
        row_text = ", ".join([f"{col}: {val}" for col, val in row.items()])
        doc.add_paragraph(row_text)

        # Agregar imágenes correspondientes
        for image in images:
            row_num, col_num, image_path = image
            if row_num == i + 1:  # Ajustar al índice del DataFrame
                doc.add_picture(image_path, width=Inches(2))

    # Agregar los gráficos al final del documento
    for chart_image in chart_images:
        doc.add_paragraph("Gráfico:")
        doc.add_picture(chart_image, width=Inches(6))

    # Guardar el documento de Word
    word_file = 'res.docx'  # Reemplaza con la ruta correcta
    doc.save(word_file)

    print(f'Datos guardados en {word_file}')